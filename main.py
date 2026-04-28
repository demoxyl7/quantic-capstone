import os
import io
import json
import time
from typing import Optional, List
import tempfile
import psycopg2
from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from pdf2docx import Converter
from docx import Document # Added for DOCX manipulation
from groq import Groq 
from dotenv import load_dotenv

load_dotenv() 

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
DATABASE_URL = os.getenv("DATABASE_URL")

if DATABASE_URL and DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# --- DATABASE LOGIC ---
def init_db():
    if not DATABASE_URL: return
    try:
        conn = psycopg2.connect(DATABASE_URL, sslmode='require')
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS usage_logs (
                id SERIAL PRIMARY KEY,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                ip_address TEXT,
                cv_preview TEXT
            );
        """)
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e: print(f"DB Error: {e}")

def log_usage(ip: str, text: str):
    if not DATABASE_URL: return
    try:
        conn = psycopg2.connect(DATABASE_URL, sslmode='require')
        cur = conn.cursor()
        cur.execute("INSERT INTO usage_logs (ip_address, cv_preview) VALUES (%s, %s)", (ip, text[:500]))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e: print(f"Log Error: {e}")

# --- DOCX MANIPULATION LOGIC ---
def apply_cv_suggestions(docx_path, suggestions):
    doc = Document(docx_path)
    applied_count = 0

    for suggestion in suggestions:
        target = suggestion.get("xml_target")
        replacement = suggestion.get("replacement")
        
        if not target or not replacement: continue

        for para in doc.paragraphs:
            # JOIN RUNS: This merges fragmented XML nodes so the target can be found
            full_text = "".join(run.text for run in para.runs)
            
            if target in full_text:
                new_text = full_text.replace(target, replacement)
                # Clear existing fragmented runs
                for run in para.runs:
                    run.text = ""
                # Put the corrected text into the first run
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)
                applied_count += 1
    
    doc.save(docx_path)
    return applied_count

# --- APP SETUP ---
app = FastAPI()

@app.on_event("startup")
async def startup_event():
    init_db()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class Experience(BaseModel):
    title: str
    company: str
    period: str
    description: str

class CVStructured(BaseModel):
    name: str = "Candidate"
    skills: list[str]
    experience: list[Experience]
    education: list[str]

class CVPersonalInfo(BaseModel):
    name: str
    email: str
    phone: str
    linkedin: str
    location: str

class CVExperienceBullet(BaseModel):
    id: str
    text: str

class CVExperience(BaseModel):
    id: str
    title: str
    company: str
    location: str
    dates: str
    bullets: list[CVExperienceBullet]

class CVEducationDetail(BaseModel):
    id: str
    text: str

class CVEducation(BaseModel):
    id: str
    degree: str
    institution: str
    dates: str
    details: list[CVEducationDetail]

class CVProject(BaseModel):
    id: str
    name: str
    description: str
    technologies: list[str]

class CVStructured(BaseModel):
    personal_info: CVPersonalInfo
    summary: str
    experience: list[CVExperience]
    projects: list[CVProject] # Added projects
    education: list[CVEducation]
    certifications: list[str]
    skills: list[str]

class JDData(BaseModel):
    role: str
    skills: list[str]
    responsibilities: list[str]
    qualifications: list[str]

class CVExtraction(BaseModel):
    cv_data: CVStructured
    jd_data: JDData

class Suggestion(BaseModel):
    id: str
    target_id: str
    type: str # summary, experience_bullet, project_description, education_detail, add_experience_bullet, add_project
    issue: str
    original_text: str
    replacement_text: str
    reason: str

class SkillGapCourse(BaseModel):
    topic: str
    description: str

class AnalysisResponse(BaseModel):
    is_cv: bool
    error_message: Optional[str] = None
    score: int
    match_status: str
    matched_skills: list[str]
    missing_skills: list[str]
    extraction: CVExtraction
    suggestions: list[Suggestion]
    skill_gap_courses: list[SkillGapCourse]

class AnalysisRequest(BaseModel):
    cv_text: str
    job_description: str

class CoverLetterRequest(BaseModel):
    cv_text: str
    job_description: str

@app.get("/")
async def root():
    return {"status": "online", "db": bool(DATABASE_URL)}

@app.post("/analyze")
async def analyze_cv(request: AnalysisRequest):
    if not client:
        raise HTTPException(status_code=500, detail="AI Client not initialized.")
        
    client_ip = client_request.headers.get("x-forwarded-for") or client_request.client.host
    log_usage(client_ip, request.cv_text)    
    
    cv_text = request.cv_text[:6000]
    jd_text = request.job_description[:6000]

    prompt = f"""
    You are an expert career consultant. Analyze the following CV and Job Description.
    
    STEP 1: Extract the CV data VERBATIM. 
    - Include Personal Info (Name, Email, Phone, LinkedIn, Location).
    - Include a Professional Summary.
    - Include ALL Experience records. Assign each record an ID like "exp_1" and each bullet point an ID like "b_1".
    - Include ALL Projects. Assign each record an ID like "proj_1".
    - Include ALL Education records. Assign each record an ID like "edu_1" and each detail bullet an ID like "ed_1".
    - Include ALL Certifications.
    - Include ALL Skills.
    
    STEP 2: Extract structured data from the JD (Role, Skills, Responsibilities, Qualifications).
    
    STEP 3: Generate a match score (0-100) and match status.
    
    STEP 4: Provide EXACTLY 5 targeted improvement suggestions. 
    - Each suggestion must reference a `target_id` from the extracted CV data.
    - `type` MUST be one of: "summary", "experience_bullet", "project_description", "education_detail", "add_experience_bullet".
    - Use "add_experience_bullet" to suggest a NEW bullet point for an experience record (target_id should be the exp_id).
    - Provide the `original_text` (empty for additions) and a `replacement_text` that better aligns with the JD.
    
    STEP 5: Suggest 3-4 course topics to bridge skill gaps.
    
    prompt = f"""
    Analyze the CV and JD. Return ONLY JSON.
    CV: {request.cv_text[:5000]}
    JD: {request.job_description[:5000]}
    
    CRITICAL: For suggestions, the 'xml_target' MUST be an EXACT, 100% case-sensitive 
    phrase from the CV text. 

    {{
      "is_cv": boolean,
      "error_message": "string",
      "score": number,
      "match_status": "string",
      "matched_skills": ["string"],
      "missing_skills": ["string"],
      "extraction": {{
        "cv_data": {{ 
          "personal_info": {{ "name": "", "email": "", "phone": "", "linkedin": "", "location": "" }},
          "summary": "",
          "experience": [
            {{ "id": "exp_1", "title": "", "company": "", "location": "", "dates": "", "bullets": [ {{ "id": "b_1", "text": "" }} ] }}
          ],
          "projects": [
            {{ "id": "proj_1", "name": "", "description": "", "technologies": [] }}
          ],
          "education": [
            {{ "id": "edu_1", "degree": "", "institution": "", "dates": "", "details": [ {{ "id": "ed_1", "text": "" }} ] }}
          ],
          "certifications": [],
          "skills": []
        }},
        "jd_data": {{ "role": "", "skills": [], "responsibilities": [], "qualifications": [] }}
      }},
      "suggestions": [
        {{ "id": "s1", "target_id": "b_1", "type": "experience_bullet", "issue": "", "original_text": "", "replacement_text": "", "reason": "" }}
      ],
      "skill_gap_courses": []
    }}
    
    IMPORTANT RULES:
    1. EXTRACT DATA VERBATIM: Do not summarize or paraphrase original CV text during extraction.
    2. NO HALLUCINATIONS: If a piece of information is missing, leave the field empty.
    3. TARGETED SUGGESTIONS: Only suggest improvements for sections that exist.
    4. NON-CV CONTENT: If the document isn't a CV, set `is_cv` to false.
    """
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a career consultant specialized in high-end recruitment analysis. Your output is always strictly structured JSON. You must be extremely literal during extraction and never hallucinate data that isn't in the provided text."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content.strip())
        
        # Guard against non-CV documents
        if result.get("is_cv") == False:
            # We still return the object but the frontend should handle the error_message
            pass
            
        return result

    except Exception as e:
        print(f"ANALYSIS ERROR: {e}")
        raise HTTPException(status_code=500, detail=f"AI Analysis failed: {str(e)}")

@app.post("/apply-suggestions")
async def apply_edits(file: UploadFile = File(...), suggestions_json: str = File(...)):
    # This endpoint allows the user to upload their DOCX and apply the AI's suggestions
    suggestions = json.loads(suggestions_json)
    
    try:
        content = await file.read()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        count = apply_cv_suggestions(tmp_path, suggestions)
        
        with open(tmp_path, "rb") as f:
            updated_data = f.read()
            
        os.remove(tmp_path)
        
        return Response(
            content=updated_data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Improved_CV.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Edit failed: {str(e)}")

# Restore the original PDF to DOCX route
@app.post("/convert-pdf-to-docx")
async def convert_pdf_to_docx(file: UploadFile = File(...)):
    try:
        content = await file.read()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pdf_file:
            pdf_file.write(content)
            pdf_path = pdf_file.name
        docx_path = pdf_path.replace(".pdf", ".docx")
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        with open(docx_path, "rb") as docx_file:
            docx_data = docx_file.read()
        os.remove(pdf_path)
        if os.path.exists(docx_path): os.remove(docx_path)
        return Response(content=docx_data, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))